from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from markdown_utils import SCREENSHOT_PLACEHOLDER_RE, cleanup_inline_markdown


SPECIAL_CENTERED_HEADINGS = {
    "摘要": "abstract_heading_cn",
    "abstract": "abstract_heading_en",
    "参考文献": "references_heading",
    "致谢": "ack_heading",
}

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}


def default_style_profile() -> dict[str, Any]:
    return {
        "styles": {
            "title": centered_style("黑体", "Times New Roman", 18, bold=True),
            "heading1": paragraph_style("黑体", "Times New Roman", 18, bold=True, space_before=10, space_after=10),
            "heading2": paragraph_style("黑体", "Times New Roman", 15, bold=True, space_before=10, space_after=10),
            "heading3": paragraph_style("黑体", "Times New Roman", 12, bold=True, space_before=10, space_after=10),
            "abstract_heading_cn": centered_style("黑体", "Times New Roman", 18, bold=True, page_break_before=True),
            "abstract_heading_en": centered_style(
                "Times New Roman",
                "Times New Roman",
                18,
                bold=True,
                page_break_before=True,
            ),
            "references_heading": centered_style("黑体", "Times New Roman", 18, bold=True, page_break_before=True),
            "ack_heading": centered_style("黑体", "Times New Roman", 18, bold=True, page_break_before=True),
            "body_cn": paragraph_style("宋体", "Times New Roman", 10.5, first_line_indent=21),
            "body_en": paragraph_style("Times New Roman", "Times New Roman", 12, first_line_indent=21),
            "keywords_cn_label": run_style("黑体", "Times New Roman", 12, bold=True),
            "keywords_cn_content": run_style("宋体", "Times New Roman", 12),
            "keywords_en_label": run_style("Times New Roman", "Times New Roman", 12, bold=True),
            "keywords_en_content": run_style("Times New Roman", "Times New Roman", 12),
            "keywords_paragraph": paragraph_style("宋体", "Times New Roman", 10.5, first_line_indent=0),
            "figure_caption": centered_style("宋体", "Times New Roman", 10.5, line_spacing=1, line_spacing_rule="single"),
            "table_caption": centered_style("宋体", "Times New Roman", 10.5, line_spacing=1, line_spacing_rule="single"),
            "table_text": centered_style("宋体", "Times New Roman", 10.5),
            "references_body": paragraph_style("宋体", "Times New Roman", 10.5, first_line_indent=-21, left_indent=21),
            "missing_asset": centered_style("楷体", "Times New Roman", 10.5),
            "code": paragraph_style("Consolas", "Consolas", 9, first_line_indent=0),
        }
    }


def paragraph_style(
    east_asia_font: str,
    latin_font: str,
    size_pt: float,
    *,
    bold: bool = False,
    alignment: str = "left",
    line_spacing: float = 1.25,
    line_spacing_rule: str = "multiple",
    first_line_indent: float = 0,
    left_indent: float = 0,
    space_before: float = 0,
    space_after: float = 0,
    page_break_before: bool = False,
) -> dict[str, Any]:
    return {
        "east_asia_font": east_asia_font,
        "latin_font": latin_font,
        "size_pt": size_pt,
        "bold": bold,
        "alignment": alignment,
        "line_spacing": line_spacing,
        "line_spacing_rule": line_spacing_rule,
        "first_line_indent_pt": first_line_indent,
        "left_indent_pt": left_indent,
        "space_before_pt": space_before,
        "space_after_pt": space_after,
        "page_break_before": page_break_before,
    }


def centered_style(
    east_asia_font: str,
    latin_font: str,
    size_pt: float,
    *,
    bold: bool = False,
    line_spacing: float = 1.25,
    line_spacing_rule: str = "multiple",
    page_break_before: bool = False,
) -> dict[str, Any]:
    return paragraph_style(
        east_asia_font,
        latin_font,
        size_pt,
        bold=bold,
        alignment="center",
        line_spacing=line_spacing,
        line_spacing_rule=line_spacing_rule,
        first_line_indent=0,
        space_before=0,
        space_after=0,
        page_break_before=page_break_before,
    )


def run_style(east_asia_font: str, latin_font: str, size_pt: float, *, bold: bool = False) -> dict[str, Any]:
    return {
        "east_asia_font": east_asia_font,
        "latin_font": latin_font,
        "size_pt": size_pt,
        "bold": bold,
    }


def infer_fonts(style: dict[str, Any], fallback: dict[str, Any]) -> dict[str, Any]:
    font = style.get("font")
    if font and "east_asia_font" not in style:
        style["east_asia_font"] = font
    if font and "latin_font" not in style:
        style["latin_font"] = font if re.search(r"[A-Za-z]", str(font)) and not re.search(r"[\u4e00-\u9fff]", str(font)) else fallback["latin_font"]
    return style


def deep_merge(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    merged = json.loads(json.dumps(base))
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = deep_merge(merged[key], value)
        else:
            merged[key] = value
    return merged


def load_style_profile(path: Path | None) -> dict[str, Any]:
    profile = default_style_profile()
    if path is None or not path.exists():
        return profile

    raw = json.loads(path.read_text(encoding="utf-8"))
    if "styles" not in raw:
        raw = {"styles": raw}
    profile = deep_merge(profile, raw)

    styles = profile["styles"]
    fallback_map = default_style_profile()["styles"]
    for name, style in list(styles.items()):
        if isinstance(style, dict):
            styles[name] = infer_fonts(style, fallback_map.get(name, fallback_map["body_cn"]))

    # Allow simplified analyzer output to drive keyword/caption styles.
    if "keywords_cn" in styles:
        styles["keywords_paragraph"] = deep_merge(styles["keywords_paragraph"], styles["keywords_cn"])
        styles["keywords_cn_label"] = deep_merge(styles["keywords_cn_label"], styles["keywords_cn"])
        styles["keywords_cn_content"] = deep_merge(styles["keywords_cn_content"], styles["keywords_cn"])
    if "keywords_en" in styles:
        styles["keywords_paragraph"] = deep_merge(styles["keywords_paragraph"], styles["keywords_en"])
        styles["keywords_en_label"] = deep_merge(styles["keywords_en_label"], styles["keywords_en"])
        styles["keywords_en_content"] = deep_merge(styles["keywords_en_content"], styles["keywords_en"])
    if "figure_caption" in styles:
        styles["figure_caption"] = infer_fonts(styles["figure_caption"], fallback_map["figure_caption"])
    if "table_caption" in styles:
        styles["table_caption"] = infer_fonts(styles["table_caption"], fallback_map["table_caption"])

    return profile


def set_run_fonts(run, east_asia_font: str, latin_font: str, size_pt: float, *, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = latin_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)


def apply_paragraph_style(paragraph, style: dict[str, Any]) -> None:
    paragraph.alignment = ALIGNMENT_MAP.get(style.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    paragraph.paragraph_format.line_spacing_rule = (
        WD_LINE_SPACING.SINGLE if style.get("line_spacing_rule") == "single" else WD_LINE_SPACING.MULTIPLE
    )
    paragraph.paragraph_format.line_spacing = style.get("line_spacing", 1.25)
    paragraph.paragraph_format.first_line_indent = Pt(style.get("first_line_indent_pt", 0))
    paragraph.paragraph_format.left_indent = Pt(style.get("left_indent_pt", 0))
    paragraph.paragraph_format.space_before = Pt(style.get("space_before_pt", 0))
    paragraph.paragraph_format.space_after = Pt(style.get("space_after_pt", 0))
    if style.get("page_break_before"):
        add_page_break_before(paragraph)
    for run in paragraph.runs:
        set_run_fonts(
            run,
            style["east_asia_font"],
            style["latin_font"],
            style["size_pt"],
            bold=style.get("bold", False),
        )


def add_page_break_before(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    page_break_before = OxmlElement("w:pageBreakBefore")
    p_pr.append(page_break_before)


def load_image_map(path: Path | None) -> dict[str, Path]:
    if path is None or not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    return {k: Path(v) for k, v in data.items()}


def add_image(doc: Document, path: Path) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(14.5))


def add_missing_asset_placeholder(doc: Document, label: str, styles: dict[str, Any]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"【待补素材：{label}】")
    apply_paragraph_style(paragraph, styles["missing_asset"])


def normalize_text(text: str) -> str:
    return cleanup_inline_markdown(text)


def add_markdown_table(doc: Document, lines: list[str], styles: dict[str, Any]) -> None:
    rows = []
    for row in lines:
        cells = [normalize_text(cell.strip()) for cell in row.strip().strip("|").split("|")]
        if all(re.fullmatch(r"[:\- ]+", cell or "") for cell in cells):
            continue
        rows.append(cells)
    if len(rows) < 2:
        return
    headers = rows[0]
    data_rows = rows[1:]
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for row_index, row in enumerate(data_rows, start=1):
        for cell_index, text in enumerate(row):
            if cell_index < len(table.rows[row_index].cells):
                table.rows[row_index].cells[cell_index].text = text
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                apply_paragraph_style(paragraph, styles["table_text"])


def apply_keyword_runs(paragraph, label: str, content: str, styles: dict[str, Any]) -> None:
    apply_paragraph_style(paragraph, styles["keywords_paragraph"])
    if label.startswith("关键词"):
        label_style = styles["keywords_cn_label"]
        content_style = styles["keywords_cn_content"]
    else:
        label_style = styles["keywords_en_label"]
        content_style = styles["keywords_en_content"]

    label_run = paragraph.add_run(label)
    set_run_fonts(
        label_run,
        label_style["east_asia_font"],
        label_style["latin_font"],
        label_style["size_pt"],
        bold=label_style.get("bold", False),
    )

    if content:
        spacer = "" if label.endswith(("：", ":")) else " "
        content_run = paragraph.add_run(f"{spacer}{normalize_text(content)}")
        set_run_fonts(
            content_run,
            content_style["east_asia_font"],
            content_style["latin_font"],
            content_style["size_pt"],
            bold=content_style.get("bold", False),
        )


def build_doc(source: Path, image_map: dict[str, Path], profile: dict[str, Any]) -> Document:
    styles = profile["styles"]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    lines = source.read_text(encoding="utf-8").splitlines()
    current_section = ""
    in_code = False
    code_lang = ""
    pending_mermaid = False
    seen_content = False
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                in_code = False
                if code_lang == "mermaid":
                    pending_mermaid = True
                code_lang = ""
            elif code_lang != "mermaid":
                paragraph = doc.add_paragraph()
                paragraph.add_run(line.rstrip())
                apply_paragraph_style(paragraph, styles["code"])
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped == "---":
            index += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_lang = stripped[3:].strip().lower()
            index += 1
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.add_run(normalize_text(stripped[2:].strip()))
            apply_paragraph_style(paragraph, styles["title"])
            seen_content = True
            index += 1
            continue

        if stripped.startswith("## "):
            text = normalize_text(stripped[3:].strip())
            normalized = text.replace(" ", "").lower()
            style_key = SPECIAL_CENTERED_HEADINGS.get(normalized, "heading1")
            paragraph = doc.add_paragraph()
            paragraph.add_run(text)
            if seen_content and style_key == "heading1":
                style = deep_merge(styles["heading1"], {"page_break_before": True})
            else:
                style = styles[style_key]
            apply_paragraph_style(paragraph, style)
            current_section = text
            seen_content = True
            index += 1
            continue

        if stripped.startswith("### "):
            paragraph = doc.add_paragraph()
            paragraph.add_run(normalize_text(stripped[4:].strip()))
            apply_paragraph_style(paragraph, styles["heading2"])
            seen_content = True
            index += 1
            continue

        if stripped.startswith("#### "):
            paragraph = doc.add_paragraph()
            paragraph.add_run(normalize_text(stripped[5:].strip()))
            apply_paragraph_style(paragraph, styles["heading3"])
            seen_content = True
            index += 1
            continue

        keyword_match = re.match(r"^(关键词[:：]|Keywords[:：])\s*(.*)$", stripped)
        if keyword_match:
            paragraph = doc.add_paragraph()
            apply_keyword_runs(paragraph, keyword_match.group(1), keyword_match.group(2), styles)
            seen_content = True
            index += 1
            continue

        image_match = SCREENSHOT_PLACEHOLDER_RE.match(stripped)
        if image_match:
            label = image_match.group(1).strip()
            image_path = image_map.get(label)
            if image_path and image_path.exists():
                add_image(doc, image_path)
            else:
                add_missing_asset_placeholder(doc, label, styles)
            seen_content = True
            index += 1
            continue

        if re.match(r"^图\s*\d+(\.\d+)?", stripped):
            caption = normalize_text(stripped)
            if pending_mermaid:
                image_path = image_map.get(caption)
                if image_path and image_path.exists():
                    add_image(doc, image_path)
                else:
                    add_missing_asset_placeholder(doc, caption, styles)
                pending_mermaid = False
            paragraph = doc.add_paragraph()
            paragraph.add_run(caption)
            apply_paragraph_style(paragraph, styles["figure_caption"])
            seen_content = True
            index += 1
            continue

        if re.match(r"^表\s*\d+(\.\d+)?", stripped):
            paragraph = doc.add_paragraph()
            paragraph.add_run(normalize_text(stripped))
            apply_paragraph_style(paragraph, styles["table_caption"])
            seen_content = True
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, table_lines, styles)
            seen_content = True
            continue

        paragraph = doc.add_paragraph()
        paragraph.add_run(normalize_text(stripped))
        if current_section.replace(" ", "") == "参考文献" and re.match(r"^\[\d+\]", stripped):
            apply_paragraph_style(paragraph, styles["references_body"])
        elif current_section.lower() == "abstract":
            apply_paragraph_style(paragraph, styles["body_en"])
        else:
            apply_paragraph_style(paragraph, styles["body_cn"])
        seen_content = True
        index += 1

    return doc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate thesis DOCX from markdown source.")
    parser.add_argument("source", type=Path)
    parser.add_argument("target", type=Path)
    parser.add_argument("legacy_image_map", nargs="?", type=Path, help="Optional image map for backward compatibility.")
    parser.add_argument("--image-map", dest="image_map", type=Path, default=None)
    parser.add_argument("--style-profile", dest="style_profile", type=Path, default=None)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    image_map_path = args.image_map or args.legacy_image_map
    profile = load_style_profile(args.style_profile)
    image_map = load_image_map(image_map_path)
    args.target.parent.mkdir(parents=True, exist_ok=True)
    document = build_doc(args.source, image_map, profile)
    document.save(args.target)
    print(args.target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
