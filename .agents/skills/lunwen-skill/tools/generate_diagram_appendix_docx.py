from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


CODE_BLOCK_RE = re.compile(r"^```(mermaid|plantuml)\s*$", re.I)
CAPTION_RE = re.compile(r"^(图\s*\d+(?:\.\d+)?\s+.+)$")


def set_run_fonts(run, east_asia_font: str, latin_font: str, size_pt: float, *, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = latin_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)


def add_paragraph(doc: Document, text: str, *, font_cn: str = "宋体", font_en: str = "Times New Roman", size: float = 10.5, bold: bool = False, center: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_fonts(run, font_cn, font_en, size, bold=bold)
    if not center:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.25


def add_code_block(doc: Document, code: str) -> None:
    for line in code.splitlines() or [""]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(line.rstrip())
        set_run_fonts(run, "Consolas", "Consolas", 9)


def extract_blocks(markdown_path: Path) -> list[dict[str, str]]:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    blocks: list[dict[str, str]] = []
    in_code = False
    lang = ""
    buffer: list[str] = []
    pending: list[dict[str, str]] = []

    for line in lines:
        stripped = line.strip()
        if in_code:
            if stripped.startswith("```"):
                pending.append({"lang": lang.lower(), "code": "\n".join(buffer)})
                in_code = False
                lang = ""
                buffer = []
            else:
                buffer.append(line.rstrip())
            continue

        code_match = CODE_BLOCK_RE.match(stripped)
        if code_match:
            in_code = True
            lang = code_match.group(1)
            buffer = []
            continue

        caption_match = CAPTION_RE.match(stripped)
        if caption_match and pending:
            item = pending.pop(0)
            item["caption"] = caption_match.group(1).strip()
            blocks.append(item)

    return blocks


def build_doc(title: str, markdown_path: Path, output_path: Path) -> Path:
    blocks = extract_blocks(markdown_path)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    add_paragraph(doc, f"{title}-附件", font_cn="黑体", font_en="Times New Roman", size=18, bold=True, center=True)
    add_paragraph(doc, "流程图、E-R 图及其他 Mermaid / PlantUML 源码附件", font_cn="黑体", font_en="Times New Roman", size=14, center=True)

    if not blocks:
        add_paragraph(doc, "正文中未提取到 Mermaid 或 PlantUML 代码块。", size=10.5)
    else:
        for index, block in enumerate(blocks, start=1):
            add_paragraph(doc, f"{index}. {block['caption']}", font_cn="黑体", font_en="Times New Roman", size=14, bold=True)
            add_paragraph(doc, f"图类型：{block['lang']}", size=10.5)
            add_code_block(doc, block["code"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate appendix DOCX containing Mermaid / PlantUML source blocks from thesis markdown.")
    parser.add_argument("title", help="Thesis title used for the appendix heading")
    parser.add_argument("markdown_file", type=Path)
    parser.add_argument("target_docx", type=Path)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    output = build_doc(args.title, args.markdown_file, args.target_docx)
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
